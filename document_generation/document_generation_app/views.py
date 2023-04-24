import os
import pymssql as pymssql
import requests
from django.core.management.commands import shell
from django.http import HttpResponse
from django.shortcuts import render, redirect
from .forms import *
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import *
from win32com.shell import shell, shellcon
from datetime import datetime
from openpyxl import *
import pytz

# import function of document generations
from document_generation_app.document_generation_functions.generation_about_arrival import Generation_about_arrival
from document_generation_app.document_generation_functions.generation_payment_order_for_advance_payment import Generate_Generation_payment_order_for_advance_payment
from document_generation_app.document_generation_functions.generation_gpc_contract import Generation_GPH_contract
from document_generation_app.document_generation_functions.generation_employment_contract import Generation_employment_contract_document
from document_generation_app.document_generation_functions.generation_removal_older import Generation_removal_older
from document_generation_app.document_generation_functions.generation_notice_conclusion import Generation_notice_conclusion
from document_generation_app.document_generation_functions.generation_termination_notice import Generation_termination_notice

path_file = shell.SHGetKnownFolderPath(shellcon.FOLDERID_Downloads)


def CompanyAPI():
    # id_company = "b5deb54b-01bd-4bc3-87d0-21359b046e2a"
    # response = requests.get("http://secretochka.ru:48910/company-service/api/v1/manager/companyes/" + id_company).json()
    # actual_address = requests.get("http://secretochka.ru:48910/company-service/api/v1/companyes/" + id_company + "/info/actualaddresses").json()
    # response["ActualAddreses"] = actual_address
    response = {
            "id": "3fa85f64-5717-4562-b3fc-2c963f66afa6",
            "inn": "7804525530",
            "kpp": "780401001",
            "name": "Капитал Кадры",
            "organizationalForm": "ООО",
            "ogrn": "1117746358608",
            "okved": "123123123",
            "legalAddress": {
                "name": "195299, Г.Санкт-Петербург, пр-кт Гражданский, д. 119 ЛИТЕР А, офис 8"
            },
            "contactInfo": {
                "id": 0,
                "name": "Олег",
                "phone": "+78989321223",
                "email": "user@example.com"
            },
            "bank": {
                "id": 0,
                "bankId": "12341234",
                "correspondentAccount": "12345123451234512345",
                "paymentAccount": "12345123451234512345"
            },
            "ActualAddreses": [{"name": "195299, Г.Санкт-Петербург, пр-кт Гражданский, д. 119 ЛИТЕР А, офис 8"}]
    }

    return response

def IndividualAPI():
    pass

def Date_conversion(date, type=None):
    arr_date = date.split('-')
    day = arr_date[0]
    arr_month = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
    month = arr_date[1]
    month = int(month)
    month_conversion = arr_month[month-1]
    year = arr_date[2]

    if type == 'quotes':
        if day[0] == '0':
            day = day[1]
        date = '\"' + day + '\" ' + month_conversion + ' ' + year + ' г.'
    elif type == 'word_month':
        if day[0] == '0':
            day = day[1]
        date = day + ' ' + month_conversion + ' ' + year + ' г.'
    else:
        date = day + '.' + arr_date[1] + '.' + year

    return date


# Create your views here.
def index(request):
    return render(request,'document_generation_app/index.html')

def Employment_contract_Document(request):
    return render(request,'document_generation_app/employment_contract.html')


def Generate_Employment_contract_Document(request):
    if request.method == 'POST':
        Generation_employment_contract_document(request)
    return redirect('employment_contract')


def GPC_Agreement(request):
    return render(request, 'document_generation_app/gpc_contract.html')


def Generate_GPC_Document(request):
    if request.method == 'POST':
        Generation_GPH_contract(request)
    return redirect('gpc_contract')



def Removal_order(request):
    return render(request, 'document_generation_app/removal_order.html')


def Generate_removal_older(request):
    if request.method == 'POST':
        Generation_removal_older(request)
    return redirect('removal_order')


def Notice_conclusion(request):
    return render(request, 'document_generation_app/notice_conclusion.html')


def Generate_Notice_conclusion(request):
    if request.method == 'POST':
        Generation_notice_conclusion(request)
    return redirect('notice_conclusion')


def Termination_noticen(request):
    return render(request, 'document_generation_app/termination_notice.html')


def Generate_Termination_notice(request):
    if request.method == 'POST':
        Generation_termination_notice(request)
    return redirect('termination_notice')


def Right_not_to_withhold_pit(request):
    return render(request, 'document_generation_app/right_not_to_withhold_pit.html')


def Generate_Right_not_to_withhold_pit(request):
    if request.method == 'POST':
        company = CompanyAPI()
        organization = company["organizationalForm"] + ' "' + company["name"] + '"'
        list_organization = list(organization.upper())
        phone = company["contactInfo"]["phone"].upper()
        list_phone = list(phone)
        inn = company['inn']
        list_inn = list(inn)
        kpp = company['kpp']
        ogrn = company['ogrn']
        list_ogrn = list(ogrn)
        paymentAccount = company['bank']['paymentAccount'].upper()
        correspondentAccount = company['bank']['correspondentAccount'].upper()
        legalAddress = company['legalAddress']["name"].upper()
        list_legalAddress = list(legalAddress)
        ActualAddreses = company['ActualAddreses'][0]["name"].upper()
        CEO = 'Сталюкова Екатерина Александровна'

        individual = 'Гайназаров Кайратбек'.upper()
        # passport_series = 'AC'.upper()
        # passport_number = '4348554'.upper()

        code = request.POST.get('code')
        list_code = list(code)


        path_file_doc = 'document_generation_app/document_templates/right_not_to_withhold_pit.xlsx'
        doc = load_workbook(path_file_doc)
        sheet = doc.active

        list_columns = ['X', 'AA', 'AD', 'AG', 'AJ',
                        'AM', 'AP', 'AS', 'AV', 'AY', 'BB', 'BH']

        row = 2
        index = 0
        for symbol in list_inn:
            for col in range(index, len(list_columns)):
                cell = list_columns[index] + str(row)
                sheet[f'{cell}'] = symbol
                index += 1
                break


        list_columns = ['X', 'AA', 'AD', 'AG', 'AJ',
                        'AM', 'AP', 'AS', 'AV']
        row = 4
        index = 0
        for symbol in list_ogrn:
            for col in range(index, len(list_columns)):
                cell = list_columns[index] + str(row)
                sheet[f'{cell}'] = symbol
                index += 1
                break


        # Input code
        sheet["CF8"] = list_code[0]
        sheet["CJ8"] = list_code[1]
        sheet["CO8"] = list_code[2]
        sheet["CT8"] = list_code[3]


        list_columns = ['B', 'C', 'E', 'F', 'G', 'I', 'L', 'N', 'P', 'S', 'U', 'W', 'Z', 'AC', 'AF', 'AI','AL',
                        'AO', 'AR', 'AU', 'AX', 'AY', 'BA', 'BC', 'BD', 'BE', 'BF', 'BG', 'BH', 'BE', 'BJ', 'BK', 'BL',
                        'BM', 'BO', 'BP', 'BR', 'BS', 'BU', 'BV', 'BW', 'BX', 'BY', 'BZ', 'CA', 'CB', 'CC', 'CE']
        row = 11
        index = 0
        for symbol in list_organization:
            for col in range(index, len(list_columns)):
                cell = list_columns[index] + str(row)
                if list_columns[index] == 'CE':
                    sheet[f'{cell}'] = symbol
                    row += 2
                    index = 0
                    break
                sheet[f'{cell}'] = symbol
                index += 1
                break


        now_date = datetime.now(pytz.timezone('UTC'))
        now_date = Date_conversion(now_date.strftime('%d-%m-%Y'))
        list_now_date = list(now_date)

        list_columns = ['U', 'W', 'Z', 'AC', 'AF',
                        'AI', 'AL', 'AO', 'AR', 'AU']
        row = 46
        index = 0
        for symbol in list_now_date:
            for col in range(index, len(list_columns)):
                cell = list_columns[index] + str(row)
                sheet[f'{cell}'] = symbol
                index += 1
                break


        global path_file
        path = path_file
        if os.path.exists(path + '/' + 'right_not_to_withhold_pit.xlsx') == False:
            doc.save(path + '/' + 'right_not_to_withhold_pit.xlsx')
        else:
            i = 1
            while True:
                if os.path.exists(path_file + '/' + f'right_not_to_withhold_pit{i}.xlsx') == False:
                    path = path_file + '/' + f'right_not_to_withhold_pit{i}.xlsx'
                    doc.save(path)
                    break
                i += 1

    return redirect('right_not_to_withhold_pit')


def About_arrival(request):
    return render(request, 'document_generation_app/about_arrival.html')


def Generate_About_arrival(request):
    if request.method == 'POST':
        Generation_about_arrival(request)

    return redirect('about_arrival')

def Payment_order_for_advance_payment(request):
    return render(request, 'document_generation_app/payment_order_for_advance_payment.html')

def Generate_Payment_order_for_advance_payment(request):
    if request.method == 'POST':
        Generate_Generation_payment_order_for_advance_payment(request)

    return redirect('payment_order_for_advance_payment')


# {
#     "id": "1fa85f64-1727-4862-b3fc-2c963f66afa4",
#     "inn": "7804525530",
#     "kpp": "780401001",
#     "name": "Капитал Кадры",
#     "organizationalForm": "ООО",
#     "ogrn": "1117746358608",
#     "okved": "123456",
#     "legalAddress": {
#         "name": "195299, Г.Санкт-Петербург, пр-кт Гражданский, д. 119 ЛИТЕР А, офис 8"
#     },
#     "contactInfo": {
#         "id": 0,
#         "name": "Олег",
#         "phone": "+78989321223",
#         "email": "user@example.com"
#     },
#     "bank": {
#         "id": 0,
#         "bankId": "123456789",
#         "correspondentAccount": "12345123451234512345",
#         "paymentAccount": "12345123451234512345"
#     },
#     "ActualAddreses": [
#         {"name": "195299, Г.Санкт-Петербург, пр-кт Гражданский, д. 119 ЛИТЕР А, офис 8"}
#     ]
# }